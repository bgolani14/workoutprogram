import streamlit as st
from docx import Document

# Initialize session state variables
if 'client_name' not in st.session_state:
    st.session_state.client_name = ""

if 'weekly_plan' not in st.session_state:
    st.session_state.weekly_plan = {day: [] for day in ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]}

if 'exercises' not in st.session_state:
    st.session_state.exercises = []

if 'week_complete' not in st.session_state:
    st.session_state.week_complete = False

# Exercise Data
exercise_data = {
    "Back": ["Lat Pulldown", "Deadrows", "Pull-Ups", "Seated Cable Rows", "Bent Over Rows", "T-Bar Rows", "Single Arm Dumbbell Rows"],
    "Chest": ["Bench Press", "Incline Bench Press", "Dumbbell Flys", "Cable Crossovers", "Push-Ups", "Decline Bench Press", "Chest Dips"],
    "Legs (Quads)": ["Squats", "Leg Press", "Lunges", "Step-Ups", "Leg Extensions"],
    "Legs (Hams and Glutes)": ["Deadlifts", "Hip Thrusts", "Romanian Deadlifts", "Hamstring Curls", "Glute Bridges"],
    "Shoulder": ["Overhead Press", "Lateral Raises", "Front Raises", "Arnold Press", "Face Pulls", "Upright Rows"],
    "Tricep": ["Tricep Pushdowns", "Overhead Tricep Extensions", "Close-Grip Bench Press", "Dips", "Skull Crushers"],
    "Bicep": ["Bicep Curls", "Hammer Curls", "Preacher Curls", "Incline Dumbbell Curls", "Concentration Curls"]
}

# Sidebar for selecting client and day
with st.sidebar:
    st.header("Client Details")
    st.session_state.client_name = st.text_input("Client Name", value=st.session_state.client_name)

    st.header("Workout Details")
    selected_day = st.selectbox("Day of Workout", ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"])
    muscle_group = st.selectbox("Muscle Group", list(exercise_data.keys()))

    # Set or update exercise count dynamically
    st.session_state.exercise_count = st.number_input(
        "Number of Exercises", min_value=1, max_value=10, step=1, value=len(st.session_state.exercises) or 1
    )

# Adjust the exercises list to match the number of exercises
if len(st.session_state.exercises) < st.session_state.exercise_count:
    for _ in range(st.session_state.exercise_count - len(st.session_state.exercises)):
        st.session_state.exercises.append({
            "name": "", "set_type": "Normal", "sets": 1, "reps": "", "weights": "", "notes": "", "details": []
        })
elif len(st.session_state.exercises) > st.session_state.exercise_count:
    st.session_state.exercises = st.session_state.exercises[:st.session_state.exercise_count]

# Main Section - Header
st.title(f"Workout Plan for {selected_day}")
st.subheader(f"Muscle Group: {muscle_group}")

# Dynamic Form for Exercises
exercise_list = exercise_data[muscle_group]

st.header(f"Workout Details ({muscle_group})")
for i in range(st.session_state.exercise_count):
    st.subheader(f"Exercise {i + 1}")

    st.session_state.exercises[i]["name"] = st.selectbox(
        f"Select Exercise {i + 1}",
        exercise_list,
        key=f"exercise_name_{selected_day}_{muscle_group}_{i}"
    )

    st.session_state.exercises[i]["sets"] = st.radio(
        f"Number of Sets for {st.session_state.exercises[i]['name']}",
        options=[1, 2, 3, 4, 5],
        horizontal=True,
        key=f"sets_{selected_day}_{muscle_group}_{i}"
    )

    st.session_state.exercises[i]["reps"] = st.text_input(
        f"Reps for {st.session_state.exercises[i]['name']}",
        key=f"reps_{selected_day}_{muscle_group}_{i}"
    )

    st.session_state.exercises[i]["weights"] = st.text_input(
        f"Weight (kg) for {st.session_state.exercises[i]['name']}",
        key=f"weights_{selected_day}_{muscle_group}_{i}"
    )

    st.session_state.exercises[i]["notes"] = st.text_area(
        f"Notes (optional) for {st.session_state.exercises[i]['name']}",
        key=f"notes_{selected_day}_{muscle_group}_{i}"
    )

# Save button
if st.button(f"Save {selected_day} Workout"):
    st.session_state.weekly_plan[selected_day] = [
        ex for ex in st.session_state.exercises if ex.get('name')
    ]
    st.success(f"{selected_day} workout saved with {len(st.session_state.weekly_plan[selected_day])} exercises.")

    # Check if all days are filled
    if all(st.session_state.weekly_plan[day] for day in ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]):
        st.session_state.week_complete = True

# Show saved exercises for the day
st.subheader(f"Current Saved Exercises for {selected_day}")
for idx, exercise in enumerate(st.session_state.weekly_plan[selected_day]):
    st.write(f"**{idx + 1}. {exercise['name']}** - {exercise['sets']} sets x {exercise['reps']} reps at {exercise['weights']} kg")
    if exercise['notes']:
        st.write(f"    _Notes:_ {exercise['notes']}")

# Generate document when all days are completed
if st.session_state.week_complete:
    st.success("All workouts saved! You can download the plan now.")

    def generate_docx():
        doc = Document()
        doc.add_heading(f'Weekly Workout Plan - {st.session_state.client_name}', level=1)

        for day, exercises in st.session_state.weekly_plan.items():
            doc.add_heading(day, level=2)
            for exercise in exercises:
                doc.add_paragraph(
                    f"{exercise['name']} - {exercise['sets']} sets x {exercise['reps']} reps at {exercise['weights']} kg"
                )
                if exercise['notes']:
                    doc.add_paragraph(f"Notes: {exercise['notes']}")

        filename = f"Workout_Plan_{st.session_state.client_name.replace(' ', '_')}.docx"
        doc.save(filename)
        return filename

    filename = generate_docx()

    with open(filename, "rb") as file:
        st.download_button(
            label="Download Weekly Workout Plan (Word)",
            data=file,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
